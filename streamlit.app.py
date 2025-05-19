import streamlit as st
import os
from openai import OpenAI
import anthropic
from io import BytesIO
import base64
from streamlit_mic_recorder import mic_recorder
from streamlit.runtime import get_instance
from streamlit.runtime.scriptrunner import get_script_run_ctx
import re
from docx import Document
import tomllib
import hmac
import warnings
import io
import logging
import uuid
import time
import threading
from datetime import datetime, timedelta


# Create a custom logger
def get_logger():
    log = logging.getLogger(__name__)
    if not log.hasHandlers():  # Avoid adding handlers multiple times
        log.setLevel(logging.INFO)
        file_handler = logging.FileHandler("log.txt", mode="a", encoding="utf-8")
        console_handler = logging.StreamHandler()
        formatter = logging.Formatter(
            "{asctime}\t{levelname}\t{module}\t{threadName}\t{funcName}\t{lineno}\n{message}",
            style="{",
        )
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        log.addHandler(file_handler)
        log.addHandler(console_handler)
    return log


if "logger" not in st.session_state:
    st.session_state.logger = get_logger()
log = st.session_state.logger

warnings.filterwarnings("ignore", category=DeprecationWarning)


def get_uuid():
    timestamp = time.time()
    id = uuid.uuid4()
    id = f"{id}-{timestamp}"
    id = base64.urlsafe_b64encode(id.encode("utf-8")).decode("utf-8")
    return id[:8]


def get_session():
    runtime = get_instance()
    ctx = get_script_run_ctx()
    session_id = ctx.session_id
    session_info = runtime._instance.get_client(session_id)
    return session_id


def elapsed(start):
    duration = time.time() - start
    duration_td = timedelta(seconds=duration)
    days = duration_td.days
    hours, remainder = divmod(duration_td.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    dur_str = ""
    if days:
        dur_str = f"{days} days "
    if hours:
        dur_str += f"{hours} hours "
    if minutes:
        dur_str += f"{minutes} minutes "
    if seconds:
        dur_str += f"{seconds} seconds"
    return dur_str


def password_entered():
    """Checks whether a password entered by the user is correct."""
    if hmac.compare_digest(st.session_state["password"], st.secrets["password"]):
        st.session_state["password_correct"] = True
        del st.session_state["password"]  # Don't store the password.
        autoplay_audio(open("assets/unlock.mp3", "rb").read())
        log.info(f"Session Start: {get_session()}")
        
        # Initialize the audio state
        if "is_speaking" not in st.session_state:
            st.session_state.is_speaking = False
            
        # Add Noa's automatic first message to welcome the student
        if len(st.session_state.messages) <= 1:  # Only system message exists
            welcome_message = "Hi there! I'm Noa Martinez, one of the clinical instructors here at Columbia. I'll be guiding you through today's simulation. We're going to practice some change management skills in a challenging setting - implementing a flu vaccination program at a county corrections facility. You'll be meeting with Sam Richards, the Operations Manager there. He's been in his position for 14 years and, between us, he's known for being pretty resistant to change. Your goal is to persuade him to support your vaccination program despite his objections. Do you have any questions before we start, or would you like to discuss your approach?"
            st.session_state.messages.append({
                "role": "assistant", 
                "content": welcome_message,
                "agent": "noa"
            })
            # The welcome audio will be played in the main loop
            st.session_state.welcome_audio_needs_playing = True
    else:
        st.session_state["password_correct"] = False


def load_settings():
    """Load settings from settings.toml file"""
    try:
        return tomllib.load(open("settings.toml", "rb"))
    except Exception as e:
        log.exception(f"Error loading settings: {e}")
        # Return minimal default settings if loading fails
        return {
            "title": "Columbia University School of Nursing: Implementing Flu Vaccination Program",
            "error_message": "An error occurred",
            "user_name": "Public Health Nurse",
            "user_avatar": "assets/User.png",
            "assistant_name": "Noa Martinez",
            "assistant_avatar": "assets/Noa.jpg",
            "intro": "Welcome to the simulation",
            "warning": "This is a simulation",
            "instruction": "You are Noa, a nursing instructor",
            "sam_instruction": "You are Sam, a corrections manager",
            "noa_instruction": "You are Noa, a nursing instructor",
            "sam": {"name": "Sam Richards", "avatar": "assets/Sam.jpg", "voice": "onyx"},
            "noa": {"name": "Noa Martinez", "avatar": "assets/Noa.jpg", "voice": "nova"},
            "parameters": {"model": "gpt-4o", "temperature": 0.7}
        }


@st.cache_data
def local_css(file_name):
    try:
        with open(file_name) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except:
        # If CSS file is missing, use basic styling
        st.markdown("""
        <style>
        body {font-family: Arial, sans-serif;}
        .meet-sam-button {
            background-color: #0e4c92;
            color: white;
            font-weight: bold;
            padding: 0.75rem 1.5rem;
            border-radius: 5px;
            text-align: center;
            margin: 20px 0;
            cursor: pointer;
        }
        .meet-sam-button:hover {
            background-color: #0a3a73;
        }
        </style>
        """, unsafe_allow_html=True)


def speech_to_text(client, audio):
    try:
        id = audio["id"]
        log.debug(f"STT: {id}")
        audio_bio = io.BytesIO(audio["bytes"])
        audio_bio.name = "audio.wav"
        transcript = client.audio.transcriptions.create(
            model="whisper-1", response_format="text", file=audio_bio
        )
        st.session_state.processed_audio = id
        return transcript
    except Exception as e:
        log.exception("")


def text_to_speech(client, text, play_immediately=True, delay_before=0, delay_after=0):
    """Enhanced text-to-speech function with timing control"""
    try:
        # Add optional delay before speaking (for sequencing)
        if delay_before > 0:
            time.sleep(delay_before)
            
        # Set speaking state to true - this lets us know audio is playing
        st.session_state.is_speaking = True
        
        log.debug(f"TTS: {text}")
        # Use the appropriate voice based on the current active agent
        if "sam_active" in st.session_state and st.session_state.sam_active:
            current_voice = "onyx"  # Sam's voice
        else:
            current_voice = "nova"  # Noa's voice
        
        response = client.audio.speech.create(
            model="tts-1",
            voice=current_voice,
            input=text,
        )
        
        # Play the audio if requested
        if play_immediately:
            autoplay_audio(response.content)
            
            # Add optional delay after speaking (for sequencing)
            if delay_after > 0:
                time.sleep(delay_after)
            
        # Return the audio content for later use
        return response.content
        
    except Exception as e:
        log.exception(f"Error in text_to_speech: {e}")
        st.session_state.is_speaking = False
        return None
    finally:
        # After audio completes or fails, set speaking state to false
        st.session_state.is_speaking = False


def play_welcome_audio():
    """Function to play the welcome message audio after UI has loaded"""
    if "welcome_audio_needs_playing" in st.session_state and st.session_state.welcome_audio_needs_playing:
        try:
            # Get the welcome message
            welcome_message = st.session_state.messages[1]["content"]
            
            # Play the audio
            speech_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            text_to_speech(speech_client, welcome_message)
            
            # Mark as played
            st.session_state.welcome_audio_needs_playing = False
        except Exception as e:
            log.exception(f"Error playing welcome audio: {e}")


def autoplay_audio(audio_data):
    """Play audio with a unique ID to prevent conflicts"""
    try:
        # Generate a unique ID for this audio element to prevent caching issues
        audio_id = f"audio_{get_uuid()}"
        b64 = base64.b64encode(audio_data).decode("utf-8")
        
        # Create the HTML for the audio element only
        md = f"""
        <audio id="{audio_id}" autoplay>
        <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
        </audio>
        """
        st.markdown(md, unsafe_allow_html=True)
        
        # Add the script separately using components.html to avoid it being displayed
        # The script will be hidden (height=0)
        st.components.v1.html(
            f"""
            <script>
                const audio = document.getElementById('{audio_id}');
                if (audio) {{
                    console.log('Audio element found and initialized');
                }}
            </script>
            """, 
            height=0
        )
        
        return True
    except Exception as e:
        log.exception(f"Error in autoplay_audio: {e}")
        return False


# Send prompt to OpenAI and get response
def stream_response_openai(client, messages):
    try:
        log.debug(f"Sending text request to OpenAI: {messages[-1]['content']}")
        
        # Get the correct system message based on which agent is active
        if "sam_active" in st.session_state and st.session_state.sam_active:
            # If Sam is active, use Sam's instruction
            instruction = st.session_state.settings.get("sam_instruction", "You are Sam, a corrections manager")
            messages_to_send = [{"role": "system", "content": instruction}] + messages[1:]
        else:
            # Otherwise use Noa's instruction
            instruction = st.session_state.settings.get("noa_instruction", "You are Noa, a nursing instructor")
            messages_to_send = [{"role": "system", "content": instruction}] + messages[1:]
        
        # Optimize for faster responses
        stream = client.chat.completions.create(
            model="gpt-4o",  # Directly specify model for reliability
            messages=messages_to_send,
            temperature=0.7,  # Direct temperature value for reliability
            stream=True,
            max_tokens=800,  # Limiting max tokens for faster responses
        )
        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content
    except Exception as e:
        log.exception("")
        yield "I'm sorry, there was an issue generating a response. Let's try again."


def show_download():
    document = create_transcript_document()
    col1, col2 = st.columns([1, 1])
    # Button to download the full conversation transcript
    with col1:
        st.download_button(
            label="📥 Download Transcript",
            data=document,
            file_name="Transcript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def create_transcript_document():
    try:
        doc = Document()
        doc.add_heading("Conversation Transcript\n", level=1)

        for message in st.session_state.messages[1:]:
            if message["role"] == "user":
                p = doc.add_paragraph()
                p.add_run("Public Health Nurse: ").bold = True
                p.add_run(message["content"])
            else:
                p = doc.add_paragraph()
                # Use the appropriate name based on which agent responded
                if "agent" in message and message["agent"] == "sam":
                    agent_name = "Sam Richards"
                else:
                    agent_name = "Noa Martinez"
                p.add_run(f"{agent_name}: ").bold = True
                p.add_run(message["content"])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        log.exception(f"Error creating transcript: {e}")
        # Return empty document if error occurs
        empty_doc = Document()
        empty_doc.add_paragraph("Error creating transcript")
        buffer = BytesIO()
        empty_doc.save(buffer)
        buffer.seek(0)
        return buffer


def stop_current_audio():
    """Function to stop any currently playing audio"""
    if "is_speaking" in st.session_state and st.session_state.is_speaking:
        # Add JavaScript to stop all audio elements using components.html
        st.components.v1.html(
            """
            <script>
                var audioElements = document.querySelectorAll('audio');
                audioElements.forEach(function(audio) {
                    audio.pause();
                    audio.currentTime = 0;
                });
            </script>
            """, 
            height=0
        )
        
        # Reset the speaking state
        st.session_state.is_speaking = False
        # Small delay to ensure audio stops
        time.sleep(0.5)


def play_sam_intro():
    """Function to play Sam's introduction audio after UI has loaded"""
    if "sam_intro_needs_playing" in st.session_state and st.session_state.sam_intro_needs_playing:
        try:
            # Find Sam's introduction message
            sam_messages = [msg for msg in st.session_state.messages 
                         if msg.get("role") == "assistant" and msg.get("agent") == "sam"]
            
            if sam_messages:
                sam_intro = sam_messages[0]["content"]
                
                # Make sure no audio is playing
                stop_current_audio()
                time.sleep(0.5)
                
                # Play Sam's introduction
                speech_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
                text_to_speech(speech_client, sam_intro)
                
                # Mark as played
                st.session_state.sam_intro_needs_playing = False
        except Exception as e:
            log.exception(f"Error playing Sam intro: {e}")


def play_debrief_intro():
    """Function to play Noa's debrief intro audio after UI has loaded"""
    if "debrief_intro_needs_playing" in st.session_state and st.session_state.debrief_intro_needs_playing:
        try:
            # Find Noa's debrief message
            noa_messages = [msg for msg in st.session_state.messages[-2:] 
                         if msg.get("role") == "assistant" and msg.get("agent") != "sam"]
            
            if noa_messages:
                debrief_intro = noa_messages[0]["content"]
                
                # Make sure no audio is playing
                stop_current_audio()
                time.sleep(0.5)
                
                # Play Noa's debrief introduction
                speech_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
                text_to_speech(speech_client, debrief_intro)
                
                # Mark as played
                st.session_state.debrief_intro_needs_playing = False
        except Exception as e:
            log.exception(f"Error playing debrief intro: {e}")


def process_user_query(text_client, speech_client, user_query):
    # Stop any currently playing audio when user inputs something new
    stop_current_audio()
    
    # Check for transition triggers
    
    # 1. Transition from Noa to Sam (pre-brief to simulation)
    if not st.session_state.sam_active and not st.session_state.debrief_active:
        # Enhanced pattern matching for readiness signals
        ready_pattern = r"ready|begin|start|yes|yeah|sure|ok|okay|yep|let'?s|meet sam|proceed|sleep|go"
        
        if re.search(ready_pattern, user_query.lower()):
            # Display the user's query first
            with st.chat_message("Public Health Nurse", avatar="assets/User.png"):
                st.markdown(user_query)
            
            # Store the user's query into the history
            st.session_state.messages.append({"role": "user", "content": user_query.strip()})
            
            # Add a state flag to track where we are in the transition process
            if "transition_step" not in st.session_state:
                st.session_state.transition_step = "noa_message"
                
                # Add only Noa's transition message first
                transition_message = "Great! I'll introduce you to Sam now. Remember to focus on addressing his specific concerns while emphasizing the benefits to his facility. Good luck!"
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": transition_message,
                    "agent": "noa"
                })
                
                # Display Noa's message and play the audio
                with st.chat_message("Noa Martinez", avatar="assets/Noa.jpg"):
                    st.markdown(transition_message)
                    
                # Play Noa's transition audio
                text_to_speech(speech_client, transition_message)
                
                # Force a rerun to ensure UI updates and audio completes
                st.rerun()
                return  # Stop further processing
                
            elif st.session_state.transition_step == "noa_message":
                # Now we've shown Noa's message, proceed to Sam
                # Make sure any audio has stopped
                stop_current_audio()
                time.sleep(0.5)  # Small pause
                
                # Update transition step
                st.session_state.transition_step = "done"
                
                # Set sam_active to True
                st.session_state.sam_active = True
                
                # Add Sam's first message
                sam_intro = "Hello, I'm Sam Richards, Operations Manager here at the County Corrections Facility. I understand you're here about some flu vaccination program? Look, I've got 500 inmates to manage, an understaffed facility, and security concerns you wouldn't believe. I'm not sure how you expect this to work in our environment. What exactly are you proposing?"
                st.session_state.messages.append({
                    "role": "assistant", 
                    "content": sam_intro,
                    "agent": "sam"
                })
                
                # Flag to play Sam's audio on next cycle
                st.session_state.sam_intro_needs_playing = True
                st.rerun()
                return  # Skip further processing
    
    # 2. Transition from Sam to Noa (simulation to debrief)
    if st.session_state.sam_active and not st.session_state.debrief_active and re.search(r"ready for feedback|end session|finish|complete|goodbye", user_query.lower()):
        # Display the user's query
        with st.chat_message("Public Health Nurse", avatar="assets/User.png"):
            st.markdown(user_query)

        # Store the user's query into the history
        st.session_state.messages.append({"role": "user", "content": user_query.strip()})
        
        # DIRECT TRANSITION to debrief mode
        # Make sure no audio is playing
        stop_current_audio()
        
        # Set state flags
        st.session_state.sam_active = False
        st.session_state.debrief_active = True
        st.session_state.end_session_button_clicked = True
        st.session_state.download_transcript = True
        
        # Add Noa's debrief introduction message
        debrief_intro = "So, how do you think that went? That wasn't easy - Sam can be quite challenging! I was really impressed with how you managed to stay calm and professional throughout the conversation, especially when he brought up some tough concerns. Let me give you some feedback on your performance."
        st.session_state.messages.append({
            "role": "assistant", 
            "content": debrief_intro,
            "agent": "noa"
        })
        
        # Set flag to play debrief audio on next cycle
        st.session_state.debrief_intro_needs_playing = True
        
        # Force a rerun to update the UI
        st.rerun()
        return
    
    # Display the user's query
    with st.chat_message("Public Health Nurse", avatar="assets/User.png"):
        st.markdown(user_query)

    # Store the user's query into the history
    st.session_state.messages.append({"role": "user", "content": user_query.strip()})

    # Stream the assistant's reply
    # Determine which agent is responding
    current_agent = "sam" if st.session_state.sam_active else "noa"
    agent_name = "Sam Richards" if st.session_state.sam_active else "Noa Martinez"
    agent_avatar = "assets/Sam.jpg" if st.session_state.sam_active else "assets/Noa.jpg"
    
    with st.chat_message(agent_name, avatar=agent_avatar):
        # Empty container to display the assistant's reply
        assistant_reply_box = st.empty()

        # A blank string to store the assistant's reply
        assistant_reply = ""

        # Iterate through the stream
        for chunk in stream_response_openai(text_client, st.session_state.messages):
            assistant_reply += chunk
            assistant_reply_box.markdown(assistant_reply)

        # Once the stream is over, update chat history with agent info
        st.session_state.messages.append(
            {"role": "assistant", "content": assistant_reply.strip(), "agent": current_agent}
        )
        
        # Play audio response AFTER the full text is shown
        text_to_speech(speech_client, assistant_reply)
        
        # Update ready_for_sam flag after Noa responds
        if not st.session_state.sam_active and not st.session_state.debrief_active:
            st.session_state.ready_for_sam = check_readiness_for_sam()


# Initialize session state
def init_session():
    if "show_intro" not in st.session_state:
        st.session_state.show_intro = True
    
    if "chat_active" not in st.session_state:
        st.session_state.chat_active = False
    
    if "sam_active" not in st.session_state:
        st.session_state.sam_active = False
    
    if "debrief_active" not in st.session_state:
        st.session_state.debrief_active = False
    
    if "processed_audio" not in st.session_state:
        st.session_state.processed_audio = None
    
    if "manual_input" not in st.session_state:
        st.session_state.manual_input = None
    
    if "end_session_button_clicked" not in st.session_state:
        st.session_state.end_session_button_clicked = False
    
    if "download_transcript" not in st.session_state:
        st.session_state.download_transcript = False
    
    if "start_time" not in st.session_state:
        st.session_state.start_time = time.time()
    
    if "settings" not in st.session_state:
        st.session_state.settings = load_settings()
    
    if "messages" not in st.session_state:
        # Initialize with a system message using Noa's instruction
        st.session_state.messages = [
            {"role": "system", "content": st.session_state.settings.get("instruction", "You are a helpful assistant")}
        ]
    
    if "ready_for_sam" not in st.session_state:
        st.session_state.ready_for_sam = False
        
    if "transition_attempted" not in st.session_state:
        st.session_state.transition_attempted = False
        
    if "is_speaking" not in st.session_state:
        st.session_state.is_speaking = False
        
    if "transition_step" not in st.session_state:
        st.session_state.transition_step = None


def setup_sidebar():
    # Simplified sidebar that doesn't rely on nested dictionary access
    if "sam_active" in st.session_state and st.session_state.sam_active:
        st.sidebar.header("Meeting with Sam Richards")
        container = st.sidebar.container(border=True)
        with container:
            st.image("assets/Sam.jpg")
            st.subheader("Name: Sam Richards")
            st.subheader("Position: Operations Manager")
            st.subheader("Years in Position: 14")
            st.subheader("Facility: County Corrections Facility")
    else:
        st.sidebar.header("Session with Noa Martinez")
        container = st.sidebar.container(border=True)
        with container:
            st.image("assets/Noa.jpg")
            st.subheader("Name: Noa Martinez")
            st.subheader("Position: Clinical Nursing Instructor")
            st.subheader("Institution: Columbia University School of Nursing")

    # Button Login
    if "password_correct" in st.session_state and st.session_state.password_correct:
        st.session_state.chat_active = True
        st.session_state.show_intro = False
    else:
        st.sidebar.header("Access Code")
        with st.sidebar.container(border=True):
            with st.form("Credentials"):
                st.text_input("Access Code", type="password", key="password")
                st.form_submit_button("Start Chat", on_click=password_entered)
            if "password_correct" in st.session_state and not st.session_state.password_correct:
                st.error("😕 Invalid Code")


def check_readiness_for_sam():
    """Check if conversation has reached a point where transition to Sam should happen"""
    # If user has indicated readiness in recent messages
    if "messages" in st.session_state and len(st.session_state.messages) >= 2:
        # Look at the last 3 user messages to check for readiness signals
        user_messages = [msg for msg in st.session_state.messages[-6:] if msg["role"] == "user"]
        if user_messages:
            for msg in user_messages:
                user_text = msg["content"].lower()
                ready_phrases = [
                    "ready to meet", "meet sam", "meet with sam", "start simulation", 
                    "begin simulation", "let's start", "i'm ready", "ready to begin",
                    "begin", "start", "yes", "sure", "okay", "ok", "yep", "yeah"
                ]
                for phrase in ready_phrases:
                    if phrase in user_text:
                        return True
                    
        # If we have enough conversation history, Noa should ask if ready
        if len(st.session_state.messages) >= 6:
            # Check if the most recent Noa message asks about readiness
            noa_messages = [msg for msg in st.session_state.messages[-3:] 
                         if msg["role"] == "assistant" and ("agent" not in msg or msg["agent"] == "noa")]
            if noa_messages:
                for msg in noa_messages:
                    noa_text = msg["content"].lower()
                    if any(phrase in noa_text for phrase in ["ready to start", "ready to meet", "would you like to meet", "shall we begin"]):
                        # Noa asked if ready; next user message could trigger transition
                        return True
            
            # If we've had a lot of back and forth with Noa, suggest transition
            if len(st.session_state.messages) >= 10:
                return True
    
    return False


def show_messages():
    for message in st.session_state.messages[1:]:
        if message["role"] == "user":
            name = "Public Health Nurse"
            avatar = "assets/User.png"
        else:
            # Determine which agent's info to use based on the message
            if "agent" in message and message["agent"] == "sam":
                name = "Sam Richards"
                avatar = "assets/Sam.jpg"
            else:
                name = "Noa Martinez"
                avatar = "assets/Noa.jpg"
                
        with st.chat_message(name, avatar=avatar):
            st.markdown(message["content"])


def handle_audio_input(client):
    with st.sidebar.container(border=True):
        audio = mic_recorder(
            start_prompt="🎙 Record",
            stop_prompt="📤 Stop",
            just_once=False,
            use_container_width=True,
            format="wav",
            key="recorder",
        )
    # Check if there is a new audio recording and it has not been processed yet
    if audio and "processed_audio" in st.session_state and audio["id"] != st.session_state.processed_audio:
        try:
            transcript = speech_to_text(client, audio)
            return transcript
        except:
            return None


def main():
    try:
        # Initialize session state first
        init_session()
        
        # Inject CSS for custom styles
        local_css("style.css")

        # Create API clients with error handling
        try:
            text_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            speech_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        except Exception as e:
            log.exception(f"Error creating API clients: {e}")
            st.error("Unable to connect to OpenAI API. Please check your API key.")
            return

        # Set page title from settings or use default
        title = st.session_state.settings.get("title", "Columbia University School of Nursing: Implementing Flu Vaccination Program")
        st.title(title)
        
        # Setup sidebar (simplified version that doesn't rely on sidebar dictionary)
        setup_sidebar()

        # Display intro text before Start Chat is pressed
        if st.session_state.show_intro:
            with st.container(border=True):
                st.markdown(st.session_state.settings.get("intro", "Welcome to the simulation"))

        # Check if chat is active
        if st.session_state.chat_active:
            # Display the messages
            show_messages()
            
            # Play welcome audio after UI has loaded (first time only)
            if "welcome_audio_needs_playing" in st.session_state and st.session_state.welcome_audio_needs_playing:
                play_welcome_audio()
                
            # Play Sam's intro after transition (if needed)
            if "sam_intro_needs_playing" in st.session_state and st.session_state.sam_intro_needs_playing:
                play_sam_intro()
                
            # Play debrief intro after transition (if needed)
            if "debrief_intro_needs_playing" in st.session_state and st.session_state.debrief_intro_needs_playing:
                play_debrief_intro()
            
            # Show "Meet with Sam Richards" button when ready
            if (not st.session_state.sam_active and 
                not st.session_state.debrief_active and 
                "ready_for_sam" in st.session_state and 
                st.session_state.ready_for_sam):
                
                # Create a prominent button to meet Sam
                if st.button("👨‍💼 I'm Ready to Meet Sam", 
                             type="primary", 
                             use_container_width=True,
                             help="Click to start your meeting with Sam Richards"):
                    # Direct transition to Sam using the simplified approach
                    with st.chat_message("Public Health Nurse", avatar="assets/User.png"):
                        st.markdown("I'm ready to meet with Sam Richards now.")
                    
                    # Simulate a user query saying they're ready
                    st.session_state.messages.append({"role": "user", "content": "I'm ready to meet with Sam Richards now."})
                    
                    # Use the new direct transition
                    process_user_query(text_client, speech_client, "I'm ready to meet with Sam Richards now.")
                    return  # Skip the rest of the function after transition

            # Check if there's a manual input and process it
            if "manual_input" in st.session_state and st.session_state.manual_input:
                user_query = st.session_state.manual_input
            else:
                # Update placeholder text based on current agent
                if "sam_active" in st.session_state and st.session_state.sam_active:
                    placeholder_text = "Chat with Sam Richards about implementing the flu vaccination program..."
                elif "debrief_active" in st.session_state and st.session_state.debrief_active:
                    placeholder_text = "Ask Noa questions about your feedback or the simulation..."
                else:
                    placeholder_text = "Chat with Noa to prepare for your meeting with Sam..."
                    
                user_query = st.chat_input(placeholder_text)
                transcript = handle_audio_input(speech_client)
                if transcript:
                    user_query = transcript

            if user_query:
                # Stop any currently playing audio before processing the new query
                stop_current_audio()
                process_user_query(text_client, speech_client, user_query)
                if "manual_input" in st.session_state and st.session_state.manual_input:
                    st.session_state.manual_input = None
                    st.rerun()

            # Handle end session button - only show during Sam conversation
            if "sam_active" in st.session_state and st.session_state.sam_active and "debrief_active" in st.session_state and not st.session_state.debrief_active:
                if st.button("End Session & Get Feedback", type="primary", use_container_width=True):
                    # Simulate a user query to end the session
                    with st.chat_message("Public Health Nurse", avatar="assets/User.png"):
                        st.markdown("Ready for feedback on my conversation with Sam.")
                    
                    st.session_state.messages.append({"role": "user", "content": "Ready for feedback on my conversation with Sam."})
                    
                    # Use the new direct transition for debrief
                    process_user_query(text_client, speech_client, "Ready for feedback on my conversation with Sam.")
                    return

            # Show the download button during debrief
            if "download_transcript" in st.session_state and st.session_state.download_transcript:
                show_download()

        # Show warning message in sidebar
        warning_msg = st.session_state.settings.get("warning", "This is a simulation")
        st.sidebar.warning(warning_msg)
        
    except Exception as e:
        id = get_uuid()
        log.exception(f"Unhandled exception: {id}")
        
        # Get error message from settings or use default
        error_msg = st.session_state.settings.get("error_message", 
                     "😞 Oops! An unexpected error occurred. Please try again. If the error persists, please contact the administrator.")
        
        st.error(f"{error_msg}\n\n**Reference id: {id}**")


if __name__ == "__main__":
    main()
